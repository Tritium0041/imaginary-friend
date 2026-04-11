"""
DOCX 文本提取器 — 使用 python-docx 从 DOCX 文件中提取文本。
"""
from __future__ import annotations

import hashlib
import logging
from dataclasses import dataclass, field
from pathlib import Path

logger = logging.getLogger(__name__)


@dataclass
class DocxDocument:
    """DOCX 文件提取结果。"""

    filename: str
    sha256: str
    paragraphs: list[str] = field(default_factory=list)

    @property
    def full_text(self) -> str:
        return "\n".join(self.paragraphs)


class DocxExtractor:
    """从 DOCX 文件中提取文本内容。"""

    def extract(self, file_path: str | Path) -> DocxDocument:
        """从文件路径提取。"""
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")

        sha256 = self._compute_sha256(file_path)
        data = file_path.read_bytes()
        return self._extract_from_bytes(data, file_path.name, sha256)

    def extract_from_bytes(self, data: bytes, filename: str) -> DocxDocument:
        """从字节流提取。"""
        sha256 = hashlib.sha256(data).hexdigest()
        return self._extract_from_bytes(data, filename, sha256)

    def _extract_from_bytes(
        self, data: bytes, filename: str, sha256: str
    ) -> DocxDocument:
        import io
        from docx import Document as DocxFile

        doc = DocxFile(io.BytesIO(data))
        paragraphs: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

        logger.info(
            "Extracted %d paragraphs from DOCX: %s", len(paragraphs), filename
        )
        return DocxDocument(filename=filename, sha256=sha256, paragraphs=paragraphs)

    @staticmethod
    def _compute_sha256(file_path: Path) -> str:
        h = hashlib.sha256()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(8192), b""):
                h.update(chunk)
        return h.hexdigest()
